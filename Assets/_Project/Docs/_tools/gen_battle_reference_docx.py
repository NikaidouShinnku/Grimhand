# -*- coding: utf-8 -*-
"""Generate 战斗逻辑及机制参考.docx — v0.9 策划案 + 代码库对照。"""
import openpyxl
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from battle_reference_v09_sections import (
    add_turn_loop_sections,
    add_presentation_sections,
    add_card_catalog,
    add_confirmed_rulings,
    patch_keyword_burn,
)

ROOT = Path(__file__).resolve().parents[1]  # Docs/
DOCS = ROOT
XLSX = Path(r"c:\Users\Kelthuzad\Desktop\Grimhand实际内容总览表v0.9.xlsx")
OUT = DOCS / "战斗逻辑及机制参考.docx"
OUT_DESKTOP = Path(r"c:\Users\Kelthuzad\Desktop\战斗逻辑及机制参考.docx")


def set_doc_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val) if val is not None else ""
    return table


def load_xlsx_overview():
    if not XLSX.exists():
        return {}
    wb = openpyxl.load_workbook(XLSX, read_only=True, data_only=True)
    data = {"sheets": wb.sheetnames}
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i > 500:
                break
            rows.append(["" if c is None else str(c).strip() for c in row])
        data[name] = rows
    return data


def build_document():
    xlsx = load_xlsx_overview()
    doc = Document()
    set_doc_font(doc)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Grimhand 战斗逻辑及机制参考")
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "v0.9 · 权威：Grimhand_Design_Doc_v0.9.md + 总览表 xlsx\n"
        "本文档 = 策划规则 + 当前代码实现对照；改卡/改战斗请先查此文档"
    )
    sr.font.size = Pt(11)
    sr.font.name = "Microsoft YaHei"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.add_page_break()

    # TOC note
    add_title(doc, "文档说明", 1)
    add_bullets(doc, [
        "权威来源：桌面 Grimhand_Design_Doc_v0.9.md（策划）+ Grimhand实际内容总览表v0.9.xlsx（数据）。",
        "本文档描述「策划应然」与「代码现状」；两者不一致时以策划案为准，代码偏差在 §十四标注。",
        "卡牌效果原文以总览表为准，不做臆测；应对标点见 v0.9 §3.4。",
        "添加新卡前：关键词 → 状态 → EffectActionType → 应对 Condition 分段 → 演出时序（§十二）。",
        "相关交接：SessionHandoff_2026-06-10_RespondCombat.md、SessionHandoff_2026-07-02_BattleReworkNotes.md",
    ])

    # ===== 1 BATTLE FLOW =====
    add_title(doc, "一、战斗整体流程", 1)
    add_para(doc, "一场战斗从 Initialize 开始，循环「抽牌 → 规划 → 速度结算 → 回合末」，直至一方全灭。")
    add_table(doc, ["阶段", "枚举名", "主要逻辑", "代码入口"], [
        ["抽牌", "Draw", "回能、回合开始状态 tick、遗物/Boss/天赋 turn-start、双方抽牌", "BattleEngine.ProcessDrawPhase"],
        ["规划", "Planning", "敌方 AI 出牌+意图、敌方目标预掷、玩家选牌/选目标/扣费", "BattleEngine.BeginPlanning + PlanningDraft"],
        ["速度结算", "SpeedResolve", "按速度轮询出牌；应对插队；逐步 ResolveStep", "BattleEngine.ResolveTurn"],
        ["回合末", "EndOfTurn", "清护甲(最终壁垒例外)、弃手、灼烧 tick、状态到期、TurnNumber++", "BattleEngine.ProcessEndOfTurn"],
        ["战斗结束", "BattleEnd", "PlayerDefeat / PlayerVictory", "BattleEngine.EvaluateOutcome"],
    ])
    add_para(doc, "回合生命周期：", bold=True)
    add_para(doc, "Initialize → StartBattle → [Draw → Planning → CommitPlan → SpeedResolve → EndOfTurnPending] → FlushPendingEndOfTurn → EndOfTurn → Draw → …")
    add_para(doc, "演出延迟：SpeedResolve 结束后设置 EndOfTurnPending=true，须等 UI 动画播完 BattleSession.FlushPendingEndOfTurn() 才执行回合末抽牌/回能。")

    add_title(doc, "1.1 规划阶段（玩家）", 2)
    add_bullets(doc, [
        "玩家从手牌选牌加入本回合出牌队列；每张牌立即扣除能量（取消选牌退款）。",
        "若卡牌需选手：弹出目标选择（CardRules.ShouldPromptForTarget）。",
        "选目标类型：DefaultEnemy/ManualSelected（敌方）、FrontAlly/BackAlly（友方）。",
        "Reach 限制：Any / FrontAndMiddle / BackOnly / MiddleAndBack（TargetReachRules）。",
        "quick_start 关键词：规划阶段可立即结算，不进入速度队列（BattleEngine.TryResolveQuickStartCard）。",
        "虚化/锁牌：CardsLockedTurnsRemaining>0 时该角色无法继续选牌；队列中已有 LockSelfCards 卡时同角色不能再排其他牌。",
        "确认规划：CommitPlan → 冻结 PresentationSnapshot → 引擎 ResolveTurn → 事件入队 → 演出播放。",
    ])

    add_title(doc, "1.2 速度结算顺序", 2)
    add_bullets(doc, [
        "SpeedResolver.BuildResolutionOrder：双方 PlayQueue 按所属战斗员拆队列。",
        "轮询制：每轮所有仍有牌的存活角色各出 1 张；同轮内按有效速度降序，同速随机。",
        "有效速度 = 基础 Speed + 状态修正（StatusRules.GetEffectiveSpeed）。",
        "RespondResolutionPlanner：敌方攻击/防御/状态步可触发玩家应对卡 → 匹配的应对步插入敌步之前。",
        "同一敌步可匹配多张应对卡，按 PlayerPlan 中的顺序依次结算。",
    ])
    add_turn_loop_sections(doc, add_title, add_para, add_bullets, add_table)

    add_title(doc, "1.3 单张牌结算管线", 2)
    add_bullets(doc, [
        "RevealIntentIfHidden：隐藏意图在轮到该牌时揭示。",
        "CardResolvedStarted 事件 → 应对分支或普通分支。",
        "应对卡：RespondEffectExecutor（条件段）+ ExecuteUnconditionalActions（句号段）+ ExecuteFailedRespondActions（失败段）。",
        "普通卡：EffectActionExecutor.ExecuteAll（条件反应 + 无条件动作）。",
        "遗物/天赋/被动卡钩子：RelicEffectRules、TalentBattleRules、PassiveCardMechanicsRules。",
        "exhaust/bonus_hand：移出战斗或回合末清除；否则进弃牌堆。",
        "CardResolvedEnded + RecordPresentationCheckpoints。",
    ])

    # ===== 2 ENERGY =====
    add_title(doc, "二、能量系统", 1)
    add_table(doc, ["概念", "行为", "代码"], [
        ["能量上限", "EnergyMax = 配置上限 + 遗物 ExtraEnergyCap（默认 8）", "BattleState / BattleConfig"],
        ["首回合", "EnergyCurrent = EnergyMax", "EnergyRules.ApplyTurnStartRegen (IsFirstPlayerTurn)"],
        ["回合开始回复", "Restore(TurnStartEnergyRegen)，默认 4，不超过上限", "EnergyRules.Restore"],
        ["获得能量", "GainTemporary：可超过上限，不改变 EnergyMax", "EnergyRules.GainTemporary"],
        ["回复能量", "Restore：封顶 EnergyMax", "EnergyRules.Restore"],
        ["下回合回复惩罚", "PendingPlayerEnergyRegenPenaltyNextTurn（如摄魂）", "EnergyRules.ApplyTurnStartRegen"],
        ["出牌费用", "TalentBattleRules.GetEffectivePlayCost；x_cost 消耗全部剩余能量", "CardPowerRules / PlanningDraft"],
        ["灵界降临", "hand_cost_zero 状态：本回合手牌 0 费", "V09NewMechanicsRules.AdjustPlayCostForHandCostZero"],
    ])
    add_para(doc, "敌方 AI 使用 TurnStartEnergyRegen 或 EnemyTurnEnergyBudget 作为选牌预算，不维护与玩家相同的 EnergyCurrent 字段。")

    # ===== 3 DECK =====
    add_title(doc, "三、抽牌 / 弃牌 / 洗牌 / 消耗", 1)
    add_table(doc, ["操作", "规则", "代码"], [
        ["抽牌", "抽牌堆空则弃牌堆洗入；玩家可跳过污染卡；手牌满则溢出直接进弃牌堆", "DeckRules.DrawCards"],
        ["洗牌", "Fisher-Yates；玩家抽牌堆顶 pin AuthorRealmStrike", "DeckRules.ShuffleDrawPile"],
        ["打出", "手牌 → 弃牌堆", "DeckRules.MovePlayedCardToDiscard"],
        ["消耗 exhaust", "从手/抽/弃移除，IsUsable=false，本场不再出现", "DeckRules.ExhaustCard"],
        ["回合末弃手", "双方手牌全弃；IsBonusHandCard 标记清除不进弃牌堆", "DeckRules.DiscardHandAtEndOfTurn"],
        ["污染 polluted", "拥有者死亡 → 该角色所有 IsUsable 卡变不可用（非 keywords 运行时判断）", "CombatantDeathRules"],
        ["诅咒 curse", "占库/手牌位但无法打出", "CardRules.IsCurseCard"],
    ])

    # ===== 4 TARGET =====
    add_title(doc, "四、目标选择与攻击距离", 1)
    add_bullets(doc, [
        "嘲讽 Taunt：自动选敌/Reach 候选时，若嘲讽者在 Reach 内则强制唯一目标（CombatMechanicsRules.FindTauntHolder）。",
        "默认敌人：PositionRules.PickDefaultTarget — 嘲讽优先，否则有效前排。",
        "有效阵型：物理 Slot 不变；GetEffectiveSlot 按存活者 Front→Middle→Back 重排 rank。",
        "Reach 随机：UsesAutoReachRoll → 在 Reach 允许槽位内随机（规划阶段敌方预掷 PrerollEnemyAutoTargets 保证预览一致）。",
        "槽位目标：EnemyFrontSlot / AllyMiddleSlot 等按有效 rank 解析。",
        "RandomEnemy / RandomAlly / RandomEnemies：随机存活目标。",
        "LastActionActor：上次行动者（应对反击/上 debuff 给攻击者）。",
        "AllEnemies：全体攻击（AOE）。",
        "应对监视：respond_status/respond_defense 需在规划时指定监视的敌方 combatantId（ResolutionTargets）。",
    ])

    # ===== 5 RESPOND =====
    add_title(doc, "五、应对机制（Respond / Parry）", 1)
    add_para(doc, "v0.9 标点规则见 §1.4。数据层 Condition 与文案对应：", bold=True)
    add_table(doc, ["文案", "含义", "数据层 Condition", "执行路径"], [
        ["逗号「，」段", "必须应对成功才触发", "LastActionAttackOnSelf 等 ≠ None", "RespondEffectExecutor.Execute"],
        ["句号「。」段", "本回合即使没应对到也触发", "None", "ExecuteUnconditionalActions"],
        ["「若应对失败…」", "仅未匹配敌步时", "RespondArmFailed", "ExecuteFailedRespondActions"],
    ])

    add_title(doc, "5.1 应对卡识别", 2)
    add_bullets(doc, [
        "RespondRules.IsRespondCard：keywords 含 parry / respond_defense / respond_status，或任意 action.Condition≠None。",
        "⚠ 已知不一致：部分 asset 使用 respond_attack，代码不认此 keyword，仅认 parry。",
    ])

    add_title(doc, "5.2 匹配与调度", 2)
    add_bullets(doc, [
        "敌方 DealDamage 步 + WouldEnemyStepAttackCombatant(应对者) → 可触发应对攻击。",
        "respond_status：敌方 Status 牌 + 监视同一 combatant。",
        "respond_defense：敌方 Defense 牌 + 监视同一 combatant。",
        "parry 无显式 condition action 时 fallback 为 LastActionAttackOnSelf。",
        "成功：ApplyConditionalEffects=true，先执行所有匹配应对步，再执行敌步。",
        "失败：无匹配敌步，ApplyConditionalEffects=false，仍执行句号段 + RespondArmFailed 段。",
        "respond_status 成功：敌方该牌加入 SuppressedEnemyCardInstanceIds，ResolveStep 时跳过效果。",
    ])

    add_title(doc, "5.3 减伤与弹反", 2)
    add_bullets(doc, [
        "GainBlockFromLastDamagePercent：登记 RespondMitigationLayer，伤害时 ApplyMitigation 按百分比减 HP 伤。",
        "ReflectLastDamageToAttacker：按预估 incoming 伤害排队 PendingParryStrikes，敌牌结算后反击。",
        "ParryImmuneAndSlowAttacker：100% 免疫 + 对攻击者施加减速（仅 RespondEffectExecutor 路径）。",
        "ApplyStatus damage_reduction：通过 IncomingDamageReductionPercent 修饰符减伤（与 Mitigation 层不同）。",
        "演出：HadRespondDefense 或 RespondMitigatedAmount>0 → 防御姿势，非受击立绘。",
    ])

    # ===== 6 DAMAGE =====
    add_title(doc, "六、伤害 / 护甲 / 治疗", 1)
    add_para(doc, "DamageRules.ApplyDamage 主要顺序（简化）：")
    add_bullets(doc, [
        "1. 守护 Guard redirect（誓死守护替队友承伤）",
        "2. 出站威力计算 + 遗物/天赋 outgoing 修正",
        "3. Boss 首击格挡",
        "4. 护甲 Block 吸收",
        "5. 受击特质、防御属性减伤、守护 50% 减伤",
        "6. 遗物 incoming、敌方防御应对武装 DefenderRespondArmRules",
        "7. 玩家应对减伤 RespondEffectExecutor.ApplyMitigation",
        "8. 虚化 Ethereal（默认 HP 伤封顶 1；巫妖天赋可改 0 并回血）",
        "9. 扣 HP、背水一战 LastStand、Boss 狂暴",
        "10. 复活祝福 ReviveBlessing（25% MaxHP）",
        "11. 死亡 → CombatantDeathRules（污染牌库）",
        "12. 受击后：不屈、战斗意志、两界行者虚化、蛇天赋清负面等",
    ])
    add_para(doc, "ApplyBlock：修饰符 → 重甲强化 +20% → 铁壁转化（天赋禁护甲时转 pending 攻击加成）→ 加 Block。")
    add_para(doc, "ApplyHeal：治疗加成、溢出转格挡、分血仪式（恶魔回血时其他友方 30%）。")

    # ===== 7 KEYWORDS =====
    add_title(doc, "七、关键词目录（KeywordCatalog）", 1)
    kw_rows = [
        ("aoe", "AOE", "对敌方全体生效（通常 AllEnemies 目标）"),
        ("poison", "中毒×层数", "回合开始每层 1 伤，无视护甲"),
        ("damage_reduction", "减伤×层数", "受击每层 -1%"),
        ("slow", "减速×层数", "每层 -1 SPD"),
        ("damage_up", "增伤×层数", "攻击牌每层 +1%"),
        ("parry", "应对攻击", "受攻击前应对效果生效"),
        ("respond_status", "应对状态", "监视目标打状态牌时触发；成功抑制该牌"),
        ("respond_defense", "应对防御", "监视目标打防御牌时触发"),
        ("armor_up", "强固×层数", "获得护甲每层 +1 flat"),
        ("vulnerable", "易伤×层数", "受击每层 +1%"),
        ("polluted", "污染", "拥有者死亡后卡不可用"),
        ("exhaust", "消耗", "使用后移出本场战斗"),
        ("burn", "灼烧×层数", "回合 **开始** tick，每层 2 伤（v0.9 策划）"),
        ("sacrifice", "献祭", "使用后自伤 SelfDamageFlat"),
        ("armor_down", "破损×层数", "获得护甲每层 -1%"),
        ("ethereal", "虚化", "本回合每次受击 HP 伤封顶 1"),
        ("weaken", "虚弱×层数", "攻击牌每层 -1%"),
    ]
    add_table(doc, ["ID", "显示名", "机制说明"], kw_rows)

    add_title(doc, "7.1 资产中存在、Catalog 未收录的关键词", 2)
    add_table(doc, ["ID", "运行时行为"], [
        ("respond_attack", "⚠ 代码不认，应改用 parry"),
        ("quick_start", "规划阶段立即结算"),
        ("x_cost", "费用=当前剩余能量"),
        ("curse", "无法打出"),
        ("token", "标记；生成时常带 exhaust"),
        ("bonus_hand", "额外手牌，回合末清除"),
        ("self_destruct", "出牌后施法者自毁"),
        ("summon", "内容标记；实际走 SummonOrGainBlock"),
        ("inherit / usable_in_constrict", "无 C# 逻辑引用"),
    ])

    # ===== 8 STATUS =====
    add_title(doc, "八、状态效果（StatusCatalog）", 1)
    add_para(doc, "通用：层数叠加；Turns 类型到期在回合末 -1；Permanent 不过期。")
    add_para(doc, "Tick 时机：", bold=True)
    add_bullets(doc, [
        "回合开始（Draw 阶段）：中毒、亡灵毒、缠绕、延迟伤害等到期伤害",
        "回合结束（EndOfTurn）：灼烧跳伤，然后持续时间递减/移除",
        "Tick 伤害直扣 HP，不走护甲/防御属性减伤管道",
    ])

    add_title(doc, "8.1 标准战斗状态", 2)
    status_rows = [
        ("poison / necrotic_poison", "中毒/亡灵毒", "永久", "回合开始 1/层"),
        ("slow", "减速", "回合", "SPD -1/层"),
        ("burn", "灼烧", "回合", "回合 **开始** 2/层（v0.9）"),
        ("attack_up / damage_up", "增伤 flat", "回合", "出站 +1/层"),
        ("attack_up_pct", "增伤%", "回合", "出站 +1%/层"),
        ("weaken / attack_down", "虚弱", "回合", "出站 -1/层"),
        ("defense_up / armor_up", "强固 flat", "回合", "获甲 +1/层"),
        ("defense_up_pct", "强固%", "回合", "获甲 +1%/层"),
        ("armor_down / defense_down_pct", "破损", "回合", "获甲 -1%/层"),
        ("vulnerable", "易伤", "回合", "受击 +1%/层"),
        ("damage_reduction", "减伤", "回合", "受击 -1%/层"),
        ("ethereal", "虚化", "回合", "受击 HP 封顶 1"),
        ("taunt", "嘲讽", "—", "强制攻击目标"),
        ("guard", "誓死守护", "—", "替队友承伤+50%减伤"),
        ("revive_blessing", "复活祝福", "永久", "致死时 25% HP 复活"),
        ("unyielding", "不屈意志", "永久", "首次 HP≤25% 回 20"),
        ("constrict", "缠绕", "回合", "回合开始 Stacks 伤；施法者锁出牌"),
    ]
    add_table(doc, ["ID", "名", "持续", "效果"], status_rows)

    add_title(doc, "8.2 v0.9 被动卡激活标记（Permanent）", 2)
    add_para(doc, "「在本场战斗中」：从该牌生效时刻起，持续到一方全灭（所有玩家角色死亡或所有敌方角色死亡），战斗结束即清除；与「永久」（跨战斗）和「X 回合」区分。")
    scope_path = ROOT / "Docs" / "_battle_scope_cards_v09.json"
    if scope_path.exists():
        scope = json.loads(scope_path.read_text(encoding="utf-8"))
        add_bullets(doc, [f"{c['cardId']}（{c['displayName']}）" for c in scope.get("cards", [])])
    add_bullets(doc, [
        "respond_stance：应对成功 +8 甲",
        "battle_will：受 HP 伤 → 永久 +5% 增伤",
        "heavy_armor：获甲 +20%",
        "final_bulwark：回合末保留 50% 护甲",
        "last_stand：HP 将归零时剩 1",
        "plague_spread：敌中毒 tick 后 30% 传染",
        "holy_infusion_pending：下一张非自身牌再执行一次",
        "rot_avatar：回合开始全体敌 +2 永久中毒",
        "hand_cost_zero：本回合手牌 0 费",
        "eternal_void：每回合 25% MaxHP 真伤 + 维持虚化",
    ])

    # ===== 9 EFFECT ACTIONS =====
    add_title(doc, "九、卡牌动作类型（EffectActionType）", 1)
    add_para(doc, "所有卡牌效果最终由 EffectActionExecutor.ExecuteOne 分发。应对卡的条件段由 RespondEffectExecutor 单独处理。")
    actions = [
        ("DealDamage", "造成伤害（支持多段/AOE/溅射/吸血/击杀回血/献祭自伤等）"),
        ("GainBlock", "获得护甲"),
        ("Heal", "治疗（含 MaxHp%）"),
        ("ApplyStatus", "施加状态"),
        ("RemoveStatus", "移除状态层数"),
        ("SwapPositionWithFrontAlly", "与前一位友方换位"),
        ("DrawCards / DrawCardsNextTurn", "当回合抽牌 / 下回合额外抽牌"),
        ("ReflectLastDamageToAttacker", "应对：反射伤害给攻击者"),
        ("GainBlockFromLastDamagePercent", "应对：登记百分比减伤"),
        ("ParryImmuneAndSlowAttacker", "应对：完全免疫+减速攻击者"),
        ("ApplyAnubisAvatar", "阿努比斯化身 + 锁出牌 2 回合"),
        ("LockRandomPlayerPlaysThisTurn", "随机玩家本回合后续出牌跳过"),
        ("ReducePlayerEnergyRegenNextTurn", "下回合能量回复 -Value"),
        ("ArmRespondDamageRedirect", "武装：下次受玩家攻击伤害×2转嫁队友"),
        ("SummonOrGainBlock", "有空位召唤否则获甲"),
        ("ConsumeBlockDealDamage", "消耗全部护甲造成伤害"),
        ("DamagePerRespondCount", "伤害=本场应对成功次数×Value"),
        ("DoubleStatusStacks", "中毒与灼烧层数翻倍"),
        ("RecycleExhaustCardsFromDiscard", "洗回消耗牌并移除 exhaust 标签"),
        ("GainEnergy", "获得临时能量（可超上限）"),
        ("GainEnergyNextTurn", "下回合开始获得临时能量（可超上限）"),
        ("ApplyConstrict", "缠绕+锁施法者出牌"),
        ("SettlePoisonAndClear", "即时结算中毒并清除"),
        ("LockSelfCards", "锁定出牌 Value 回合"),
        ("RandomSnakeGodEffect", "蛇神回应三随机效果之一"),
        ("SealNextEnemyCard / RevealEnemyIntent", "占位/TODO"),
    ]
    add_table(doc, ["类型", "说明"], actions)

    # ===== 10 COMPLEX CARDS =====
    add_title(doc, "十、复杂卡牌与特殊互动", 1)
    add_para(doc, "以下卡牌涉及多段 action、应对、或跨系统钩子，添加类似卡时请对照。")

    complex_cards = [
        ("铁壁弹反 w_iron_parry", "parry", "30%减伤+100%反射；弹反段：敌攻后应对者 Attack 立绘反击（§12.3）"),
        ("见招拆招 w_parry_counter", "parry", "ParryImmuneAndSlowAttacker：100%免疫+2层减速"),
        ("终焉守护 m_final_guard", "parry", "逗号段：50%减伤+OnFinalGuardResponded清空能量+禁回能；句号段：GainBlock 8甲"),
        ("怨气护体 m_grudge_guard", "parry", "逗号段：90%减伤；失败段 RespondArmFailed：自身中毒×5"),
        ("不动如山 m_golem_unmovable", "parry", "逗号段：90%减伤；句号段：defense_up×5"),
        ("叹息之墙 l_wall_of_sighs", "respond_attack⚠", "80%减伤(ApplyStatus)+RandomAlly 虚化；应使用 parry keyword"),
        ("剑柄猛击 w_pommel_strike", "respond_status", "成功则抑制敌方状态牌 + 6伤"),
        ("蛇信感知 v_tongue_sense", "respond_status", "监视目标；成功对 LastActionActor 上毒"),
        ("无尽血刃", "—", "PassiveCardMechanicsRules：每次使用伤害倍率×2；出牌前献祭25%HP"),
        ("终焉魂缚 m_final_bind", "—", "敌有毒+减速 → 30层毒（PrepareGargoyleSunder 类逻辑）"),
        ("灵界降临", "hand_cost_zero", "本回合手牌 0 费；V09NewMechanicsRules"),
        ("虚化形态", "LockSelfCards", "锁定自身出牌；CardLockRules 跳过+返还能量"),
        ("沙矛重塑", "exhaust 触发", "队友消耗牌时随机打敌"),
        ("战术大师终结 w_tactician_finisher", "—", "DamagePerRespondCount：应对成功次数×伤害"),
    ]
    add_table(doc, ["卡牌", "关键词", "运行机制"], complex_cards)

    # ===== 11 ENEMY INTENT =====
    add_title(doc, "十一、敌方意图", 1)
    add_bullets(doc, [
        "EnemyTurnPlanner.PrepareEnemyTurn：贪心选手牌，生成 EnemyIntentSlot 列表。",
        "隐藏规则：≥3 张时 index 1 必藏；≥4 张时 index 2 有 50% 隐藏。",
        "SortIntentsByResolutionSpeed：与 SpeedResolver 同序展示。",
        "轮到该牌时 RevealIntentIfHidden → EnemyIntentPrepared 事件。",
        "PrerollEnemyAutoTargets：保证意图预览目标与结算一致。",
    ])

    # ===== 12 PRESENTATION =====
    add_title(doc, "十二、战斗演出与 UI", 1)
    add_presentation_sections(doc, add_title, add_para, add_bullets, add_table)

    add_title(doc, "12.6 立绘 Pose 速查", 2)
    add_table(doc, ["Pose", "用途", "资源字段"], [
        ("Idle", "规划阶段 idle 循环", "IdlePortrait / GIF 帧"),
        ("Attack", "攻击/状态牌出牌", "AttackPortrait"),
        ("Defense", "防御牌 / 应对成功受击", "DefensePortrait"),
        ("Hit", "受击（非应对成功）", "HitPortrait"),
        ("Death", "死亡", "DeathPortrait + DeadTint"),
    ])
    add_bullets(doc, [
        "出牌段：PortraitPoseChanged → MoveToCenter(仅X) → ShowPose → 效果事件 → PortraitIdleRestored → ReturnHome。",
        "CardType→Pose：Attack/Status→Attack；Defense→Defense。",
        "应对成功受击：Defense pose + 飘字/闪白，不切换 Hit pose。",
        "MoveToCenter 只改 X，Y/Z 保持 home 位置。",
    ])

    add_title(doc, "12.7 动画时长常量", 2)
    add_table(doc, ["常量", "值", "位置"], [
        ("MoveDuration", "0.35s", "CombatantPortraitView"),
        ("HitFlashDuration", "1.0s", "CombatantPortraitView"),
        ("DefenseReactDuration", "1.0s", "BattlePortraitDirector"),
        ("AttackWindUpDuration", "0.18s", "BattlePortraitDirector"),
        ("IdleFrameInterval", "0.13s（不受倍速影响）", "CombatantPortraitView"),
        ("ParryCounterDuration", "0.85s", "弹反 Attack 保持（PlayParryCounterAttack 备用）"),
        ("FastMultiplier", "2×", "BattlePresentationSpeed"),
    ])

    add_title(doc, "12.8 UI 布局", 2)
    add_table(doc, ["元素", "位置/尺寸", "文件"], [
        ("统一脚线 Y", "0.12 锚点", "CombatantSlotView"),
        ("立绘顶锚", "0.88", "CombatantSlotView"),
        ("玩家立绘缩放", "2.28", "CombatantSlotView"),
        ("敌人立绘缩放", "1.28（Boss 2.35）", "CombatantSlotView"),
        ("HP 条", "148×32，FootStatusRoot 内 Y=0", "UnitStatsRowView"),
        ("护甲 chip", "64 宽；ironWallPending 优先显示 +攻击", "UnitStatsRowView"),
        ("脚边状态图标", "Icon 28px，MaxRowWidth 144，stats 下方 -2px", "CombatantFootStatusIconsView"),
        ("伤害飘字", "底部 offset (0,-6)，字号 34", "CombatantPortraitView"),
        ("活动牌面", "HandArea 中央 BattleActiveCardBanner", "BattleScreenView"),
    ])

    # ===== 13 HOOKS =====
    add_title(doc, "十三、规则钩子（添加卡时常触及）", 1)
    add_table(doc, ["类", "文件", "典型钩子"], [
        ("天赋", "TalentBattleRules.cs", "改费、改伤、应对后 buff、虚化、中毒修正"),
        ("被动卡", "PassiveCardMechanicsRules.cs", "无尽血刃、天神下凡、终焉守护、背水一战等"),
        ("遗物", "RelicBattleRules / RelicEffectRules", "出站/incoming 修正、回合 start/end"),
        ("v0.9 机制", "V09NewMechanicsRules.cs", "缠绕、虚化计数、手牌 shuffle、蛇/巫妖 BOSS"),
        ("Boss", "BossTraitRules.cs", "骨工坊召唤、幽灵女王狂暴"),
        ("小怪特质", "MinionTraitRules.cs", "鼠群分裂、蜘蛛毒易伤、回合末护甲保留"),
        ("死亡", "CombatantDeathRules.cs", "污染牌库、有效阵型刷新"),
        ("召唤", "SummonRules.cs", "SpawnFromTemplate、自爆、bonus hand"),
    ])

    # ===== 14 KNOWN ISSUES =====
    add_title(doc, "十四、代码与策划已知偏差（改卡前必读）", 1)
    add_bullets(doc, [
        "respond_attack vs parry：总览表写【应对攻击】，资产 keyword 应统一为 parry。",
        "蛛网包裹：策划=下回合无法使用攻击牌；当前 asset 用 LockSelfCards 锁全部牌，需改 LockAttackCards。",
        "战术大师终结技：策划=远征累计应对次数；代码=单场 RespondSuccessCount。",
        "灼烧 tick：策划=回合开始；代码 StatusCatalog 仍为 TurnEndDamage。",
        "灵能体 / 灵界封印 / 恐惧低语：占位或 TODO，见 §十八。",
        "应对卡 Condition 须手配 asset，标点规则不自动解析。",
        "Generator 与 .asset 可能 drift，改卡需 xlsx → Catalog → asset 同步。",
    ])

    add_title(doc, "十四附、近期修复与机制裁定（2026-07）", 2)
    add_table(doc, ["主题", "问题", "裁定 / 修复", "影响"], [
        ["【快速启动】",
         "repair 脚本未映射 keyword，11 张卡 Keywords 被清空，规划阶段无法立即打出。",
         "extract_keywords 补 quick_start；repair_cards_by_master 描述兜底；重建 11 张卡 asset。",
         "CardView / BattleEngine.TryResolveQuickStartCard 均依赖 Keywords.Contains(\"quick_start\")。"],
        ["剑刃风暴",
         "asset Target=DefaultEnemy 导致需选手动目标；与「随机敌人×5」不符。",
         "Target=RandomEnemy(13)+HitCount=5；CardRules 对 RandomEnemy 不要求 pick。",
         "batch 行为测 w_blade_storm；repair CARD_ACTION_OVERRIDES 已固化。"],
        ["嘲讽挑衅",
         "asset 仅有 GainBlock 无 taunt 状态；敌意图在规划期预掷，不受 PickDefaultTarget 影响。",
         "ApplyStatus taunt + 按防御属性缩放护甲；ResolveTarget 优先 taunt；施加后 RefreshEnemyResolutionTargetsForTaunt。",
         "意图预览 PredictIntentTarget 同步；Reach 须能打到嘲讽者站位。"],
        ["护甲 UI",
         "应对成功等有 Block 但脚边/HP 行护甲 chip 不显示；演出末 ClearAllBlock 清空 snapshot。",
         "PresentationSnapshot.SyncBlockFromLive 替代 ClearAllBlock；Block 与脚边 status 图标分离（UnitStatsRowView）。"],
        ["遗物图标",
         "战后奖励显示占位色块+首字（如「太」），RelicVisualCatalog 未注入 UI。",
         "BattleScreenController.EnsureCatalogReferences 自动加载 RelicVisualCatalog_Demo；Entries 绑定 sun_pyramid 等。",
         "ExpeditionPostBattleOverlayView / Shop / Inventory 均走 _relicCatalog.GetIcon。"],
        ["卡牌 batch 测",
         "238 张行为断言 CardV09BehaviorBatchRunner；run_card_tests.ps1 一键跑。",
         "新卡改 asset 后必须跑绿；strict 静态审计与行为测分离见 CARD_BEHAVIOR_TEST.md。",
         "_card_behavior_verified.json 记录通过集。"],
        ["手牌选中滚动",
         "选牌后 HandScroll 跳回最左，打断选牌体验。",
         "HandPanelView 在 Layout 重建前后保存/恢复 horizontalNormalizedPosition；仅手牌张数变化时重置。",
         "选中仅高亮，不强制滚动。"],
        ["黑暗之雾 / 条件加伤 asset",
         "rebuild 丢失 Bonus* / ApplyStatus / 特殊 Type，导致虚弱、HP 条件加伤、护盾猛击等失效。",
         "补 d_dark_mist 第三段 weaken；w_fatal_strike / w_power_cleave / w_shield_slam(ConsumeBlockDealDamage) / w_burning_fury(DealDamageScaledByActorHpLoss) 字段；repair CARD_ID_OVERRIDES 固化。",
         "CombatMechanicsRules.ComputeConditionalDamageBonus 依赖 asset 字段。"],
        ["祈祷祝福 / 快速启动",
         "p_bless 用 AllyFrontSlot 自动选前排；先发制人/复活祝福需先选目标再立即结算。",
         "p_bless Target=FrontAlly；PlanningDraft 待结算 quick_start + BattleEngine.TryResolveQuickStartCard 先移出手牌再 Execute。",
         "AssignTarget → TryConsumePendingQuickStart。"],
        ["太阳之怒稀有度",
         "策划绿卡 asset Rarity=Epic 导致 UI 紫色框。",
         "p_solar_wrath Rarity=Common(0)。",
         "CardVisualResolver 按 Rarity 着色。"],
        ["背水一战",
         "LastStand 状态 asset Duration=-1；增伤未进修饰符。",
         "Duration=2；CombatModifierRules 识别 last_stand +20% OutgoingDamagePercentBonus。",
         "StatusCatalog.LastStand。"],
        ["嗜血抓挠",
         "后半段「下次攻击+3」无实现。",
         "EffectActionExecutor.TryGrantBloodScratchNextAttack；CombatantState.NextAttackFlatBonus；攻击后 MarkFirstAttackConsumed 清零。",
         "g_blood_scratch DefinitionId 钩子。"],
        ["战斗阶段抽牌保留",
         "非 Planning 阶段 DrawCards 下回合被回合末弃牌。",
         "DeckRules.TryAddToHand 标记 RetainInHandOverTurnEnd；DiscardHandAtEndOfTurn 跳过弃置且不移出手牌。",
         "快速启动抽牌：先 Remove 该牌再 Draw，7/8 抽 2 → 8/8。"],
        ["见招拆招应对",
         "EnemyStepHasAttack 仅认 DealDamage，敌人 ConsumeBlockDealDamage（护盾猛击类）不触发应对匹配。",
         "RespondTriggerMatcher.ActionDealsDamage 含 ConsumeBlockDealDamage / DealDamageScaledByActorHpLoss 等。",
         "见招拆招 ParryImmuneAndSlowAttacker + Condition=Attack。"],
        ["敌人 AI Reach",
         "对应站位全灭后仍规划只打中/后排的牌。",
         "EnemyTurnPlanner 过滤 CardRules.CardHasPlayableTargets。",
         "Reach 与 FormationSlot 联动。"],
        ["无视 N% 护甲",
         "此前 ignoreDefPercent>0 时不扣真实护甲，与「有效护甲折算后分流伤害」不符。",
         "ComputeEffectiveBlock 折算有效护甲；blocked 仍从 recipient.Block 扣除；HP 受到 afterBlock。",
         "例：10 护甲 + 50% 无视 + 10 伤 → 护甲 5，HP -5。"],
        ["神圣灌注费用",
         "x_cost 导致费用异常；应先 0 费打出，下一张选牌 +1 费。",
         "移除 x_cost；PlanningDraft 检测队列中上一张为 holy_infusion 则 GetPlayCost +1。",
         "结算仍重复下一张（holy_infusion_pending）。"],
        ["沙矛重塑",
         "误实现为被动触发或 BattleConfig 字段；应为打出时按远征消耗牌计数重复 4 伤。",
         "RunModifiers.SandSpearExhaustCardsPlayed 远征累计；打出时 4 伤×计数、随机敌人；每次消耗 +1。",
         "跨场经 ExpeditionRunState.V09SandSpearExhaustCardsPlayed 继承。"],
        ["角色悬停状态框",
         "仅有名字×层数，无描述。",
         "CombatantDetailPopupView 旁侧 StatusPopup；FormatStatusTooltipDescriptions。",
         "含 Anubis/血怒/血祭坛增伤等。"],
        ["瘟疫蔓延",
         "30% 传染、半层、相邻敌人、持续时间需与描述一致。",
         "TryTriggerPlagueSpreadOnPoisonTick：中毒 tick 后 30% 随机相邻敌人，层数=max(1,⌊原层/2⌋)，继承 RemainingTurns。",
         "需玩家侧持有 plague_spread 状态。"],
        ["死亡角色卡牌",
         "污染/绑定死亡角色的牌仍可选或归属到其他存活同职业。",
         "GetOwnerCombatantId：OwnerCombatantId 绑定死亡返回 null；IsUsable=false 不 fallback。",
         "PlanningDraft.TryGetSelectableCard 拒绝 owner 死亡。"],
        ["快速启动胜利",
         "快速启动击杀全敌需等回合结束才判胜。",
         "TryResolveQuickStartCard 末尾 EvaluateOutcome；敌人全灭即 BattleEnded Victory。",
         "不等 ProcessEndOfTurn。"],
        ["血族传承",
         "加 MaxHp 同时补满当前 HP；UI 不刷新上限。",
         "StatusRules 对 bloodline_legacy 跳过 Hp 补满；StatusApplied 触发 Refresh。",
         "10/20 → 10/30。"],
        ["血祭坛",
         "献祭减耗/增伤不触发；增伤误显示为获得护甲。",
         "AdjustSacrificeSelfDamage 对 char_ranger；OnSacrificeCardResolved 叠 SacrificeAttackStacks；反馈改 StatusApplied。",
         "暗黑献祭不再弹出护甲获得动画。"],
        ["浪潮冲锋",
         "asset 误用 BonusIfTargetHpBelowFlat=8；YAML 重复 action 块。",
         "BonusIfActorFasterThanAllEnemiesFlat=8；IsActorFasterThanAllEnemies 比较有效速度。",
         "12 伤；快于所有敌人 +8。"],
        ["手牌费用 UI",
         "HandPanelView 调用不存在的 draft.GetPlayCost(state,card)。",
         "PlanningDraft.GetPlayCost 改为 public；含神圣灌注 +1 surcharge。",
         "能量不足时卡牌不可选。"],
    ])

    add_title(doc, "十五附、2026-07-03 开发日志摘要", 2)
    add_para(doc, "完整版见 Assets/_Project/Docs/SessionSummary_2026-07-03.md")
    add_bullets(doc, [
        "无视 N% 护甲：有效护甲折算后同时扣 Block 与 HP（例：10 护甲 + 50% 无视 + 10 伤 → 护甲 5、HP 15）。",
        "神圣灌注：0 费；队列中下一张 +1 费；PlanningDraft.GetPlayCost public。",
        "沙矛重塑：远征累计消耗牌计数；打出时 4 伤×计数、随机敌人；3 费紫框攻击牌；非被动触发。",
        "UI：状态框紧贴主框、增伤等含机制描述；手牌费用含神圣灌注 surcharge。",
        "血族传承 / 血祭坛 / 瘟疫蔓延 / 快速启动胜 / 死亡角色牌 / 浪潮冲锋：见 §十四附 同表。",
        "测试：238/238 行为测通过；新增无视护甲、血族传承、沙矛、浪潮冲锋单元测试。",
        "遗留：strict 23 张静态 pending；PASSIVE 卡需手测（见 _card_honest_review_v09.md）。",
    ])

    # ===== 17-18 CARDS + PENDING =====
    add_card_catalog(doc, add_title, add_para, add_bullets, add_table)
    add_confirmed_rulings(doc, add_title, add_para, add_bullets, add_table)

    # ===== 15 XLSX =====
    add_title(doc, "十五、总览表 v0.9 摘要", 1)
    if not xlsx:
        add_para(doc, "未找到总览表 xlsx，跳过。")
    else:
        add_para(doc, f"来源：{XLSX}")
        add_para(doc, f"工作表：{', '.join(xlsx.get('sheets', []))}")
        for sheet in xlsx.get("sheets", []):
            rows = xlsx.get(sheet, [])
            if not rows:
                continue
            add_title(doc, f"15.{xlsx['sheets'].index(sheet)+1} 「{sheet}」", 2)
            # header
            header = rows[0] if rows else []
            max_cols = min(8, max(len(r) for r in rows[:1]) if rows else 0)
            if max_cols == 0:
                continue
            header = header[:max_cols]
            body = []
            for r in rows[1:201]:  # cap rows per sheet
                if not any(c for c in r):
                    continue
                body.append((r + [""] * max_cols)[:max_cols])
            if body:
                add_table(doc, header, body)
            if len(rows) > 201:
                add_para(doc, f"（仅展示前 200 行，共 {len(rows)} 行）")

    # ===== 16 FILE INDEX =====
    add_title(doc, "十六、关键代码文件索引", 1)
    files = [
        "Scripts/Battle/BattleEngine.cs — 战斗主循环",
        "Scripts/Battle/Planning/PlanningDraft.cs — 规划选牌",
        "Scripts/Battle/Rules/SpeedResolver.cs — 速度排序",
        "Scripts/Battle/Reactions/* — 应对全套",
        "Scripts/Battle/Effects/EffectActionExecutor.cs — 动作执行",
        "Scripts/Battle/Effects/DamageRules.cs — 伤害/护甲/治疗",
        "Scripts/Battle/Effects/TargetRules.cs — 目标解析",
        "Scripts/Battle/Rules/DeckRules.cs — 牌库",
        "Scripts/Battle/Rules/EnergyRules.cs — 能量",
        "Scripts/Battle/Rules/StatusRules.cs — 状态",
        "Scripts/Battle/Rules/KeywordCatalog.cs — 关键词",
        "Scripts/Battle/Status/StatusCatalog.cs — 状态定义",
        "Scripts/Battle/Model/EffectEnums.cs — 动作/目标/条件枚举",
        "Scripts/Content/CardDescriptionCatalog.cs — 卡牌描述",
        "Scripts/Presentation/Battle/BattlePortraitDirector.cs — 演出",
        "Scripts/Presentation/Battle/PresentationSnapshot.cs — 演出快照",
        "Data/Cards/*.asset — 卡牌数据",
    ]
    add_bullets(doc, files)

    matrix_md = DOCS / "_card_impl_matrix_v09.md"
    if matrix_md.exists():
        add_title(doc, "十九、全卡实现矩阵", 1)
        add_para(doc, "来源：_card_master_v09.json + _card_verification_master.json（audit 0 fail 时全为 OK）。")
        rows = []
        for line in matrix_md.read_text(encoding="utf-8").splitlines():
            if not line.startswith("| `"):
                continue
            parts = [p.strip() for p in line.strip("|").split("|")]
            if len(parts) >= 6:
                rows.append([
                    parts[0].strip("`"),
                    parts[1],
                    parts[2],
                    parts[3],
                    parts[4],
                    parts[5],
                ])
        if rows:
            add_table(doc, ["cardId", "名称", "类别", "战斗位置", "钩子", "核对状态"], rows[:250])
            if len(rows) > 250:
                add_para(doc, f"（矩阵共 {len(rows)} 行，docx 展示前 250 行；完整见 _card_impl_matrix_v09.md）")

    add_title(doc, "附录：持续更新说明", 1)
    add_bullets(doc, [
        "添加新关键词：更新 KeywordCatalog.cs + 本文档第七章。",
        "添加新状态：更新 StatusCatalog.cs + 本文档第八章 + StatusRules tick 逻辑。",
        "添加新 EffectActionType：更新 EffectEnums.cs + EffectActionExecutor + 本文档第九章。",
        "添加应对卡：对照第五章标点规则；asset 用 parry 而非 respond_attack。",
        "添加复杂互动卡：在第十章追加一行，并写清触发文件/钩子。",
        "改 UI 布局：同步更新第十二章常量表。",
    ])

    saved_any = False
    for out_path in (OUT, OUT_DESKTOP):
        try:
            doc.save(str(out_path))
            saved_any = True
        except PermissionError:
            continue
    if not saved_any:
        raise SystemExit("Could not save docx")
    print("OK", saved_any)


if __name__ == "__main__":
    build_document()
