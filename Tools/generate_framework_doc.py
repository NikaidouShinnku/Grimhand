# -*- coding: utf-8 -*-
"""Generate Grimhand framework architecture docx for review."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

OUTPUT = r"c:\Users\Kelthuzad\Desktop\The Grimhands Asset\Grimhand_Framework_Architecture_v1.2.docx"


def set_cell_shading(cell, color_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.63 * (level + 1))
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "2F5496")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("GRIMHAND / 暗手牌")
    t_run.bold = True
    t_run.font.size = Pt(22)
    t_run.font.name = "Microsoft YaHei"
    t_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run("技术框架设计文档 v1.2")
    s_run.font.size = Pt(16)
    s_run.font.name = "Microsoft YaHei"
    s_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m_run = meta.add_run(f"{date.today().isoformat()}  |  基于 GDD v0.4  |  供架构审核")
    m_run.font.size = Pt(10)
    m_run.font.color.rgb = RGBColor(100, 100, 100)
    m_run.font.name = "Microsoft YaHei"
    m_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.add_paragraph()

    add_para(
        doc,
        "本文档描述 Grimhand 客户端的技术框架方案，目标是在 PvE 最小可玩 Demo 阶段即建立可扩展骨架，"
        "避免后续 Roguelike 远征、养成、PvP、交易等系统导致战斗内核返工。",
    )
    add_para(doc, "文档状态：v1.2 已纳入审核修订（含 Planning 选牌即扣费/可取消返还）。审核通过后可启动战斗 Demo。", italic=True)

    # Ch1
    add_heading(doc, "第一章：文档目的与范围", 1)
    add_heading(doc, "1.1 目的", 2)
    add_bullet(doc, "将 GDD v0.4 中的核心战斗机制落实为可执行的软件架构。")
    add_bullet(doc, "明确「现在必须做」与「现在只留接口」的边界，控制 Demo 范围。")
    add_bullet(doc, "为团队审核提供统一的术语、分层、模块划分与验收标准。")

    add_heading(doc, "1.2 范围", 2)
    add_table(
        doc,
        ["范围", "说明"],
        [
            ["包含", "战斗模拟内核、数据驱动卡牌效果、事件流、PvE 意图接口、程序集划分、Demo 切片、测试策略"],
            ["暂不实现", "远征地图、遗物、拍卖行、PvP 联网、立绘产线（仅预留事件字段）"],
            ["对齐 GDD", "同时出牌、速度交错结算、应对机制、3 槽位、混合牌堆、能量分配、死亡污染"],
        ],
    )

    add_heading(doc, "1.3 当前工程状态", 2)
    add_para(doc, "Unity 工程为 URP 2D 空模板，尚无游戏脚本。本框架从绿场开始搭建，与 GDD 第十章优先级一致。")

    add_heading(doc, "1.4 v1.1 审核修订摘要（已确认）", 2)
    add_table(
        doc,
        ["#", "规则", "要点"],
        [
            ["1", "能量", "上限 8；战斗开始第 1 回合即为 8；未用完的能量保留；每回合开始回复 3（不超过上限）"],
            ["2", "手牌与牌堆", "手牌上限 8；每回合抽 5；溢出洗入弃牌堆；回合结束手牌全弃；打出牌入弃牌堆；抽牌堆空则弃牌堆洗回"],
            ["3", "同速", "速度相同则随机决定先手（使用战斗 RNG，可复现）"],
            ["4", "立绘", "攻击/防御/状态三类对应立绘；打出时切换，该牌结算完毕回 Idle；另有 Idle 与 Death"],
            ["5", "出牌交互", "手牌点选排序；选牌即扣能量、取消即返还；能量不足不可选；点「出牌」仅确认提交"],
        ],
    )

    # Ch2
    add_heading(doc, "第二章：架构总览", 1)
    add_heading(doc, "2.1 核心原则", 2)
    add_bullet(doc, "模拟与表现分离：战斗逻辑为纯 C#，不引用 Unity API。")
    add_bullet(doc, "数据驱动内容：卡牌/角色/敌人通过 ScriptableObject + 可组合 EffectSpec 定义。")
    add_bullet(doc, "事件驱动表现：状态变更通过 BattleEvent 流对外公布，UI/动画只消费事件。")
    add_bullet(doc, "可测试优先：速度结算等核心规则必须有单元测试，不依赖 Play Mode。")
    add_bullet(doc, "PvE 先行、接口预留：PvP / 远征 / 养成通过输入提供者与配置修饰符接入。")

    add_heading(doc, "2.2 分层结构", 2)
    add_table(
        doc,
        ["层级", "程序集（建议）", "职责"],
        [
            ["领域层", "Grimhand.Battle / Grimhand.Core", "战斗引擎、规则、效果执行、RNG、事件"],
            ["数据层", "Grimhand.Content", "ScriptableObject 定义、效果描述数据"],
            ["表现层", "Grimhand.Presentation", "MonoBehaviour、UI、立绘、音效、输入"],
            ["应用层", "（Presentation 内控制器）", "BattleController、场景流程"],
            ["远征层（后期）", "Grimhand.Expedition", "节点图、局内成长、难度修饰"],
            ["元游戏层（后期）", "Grimhand.Meta", "局外经验、天赋、存档、交易"],
            ["测试", "Grimhand.Tests", "EditMode / .NET 单元测试"],
        ],
    )

    add_para(doc, "数据流（简述）：", bold=True)
    add_para(doc, "玩家输入 → BattleController → BattleEngine（提交 BattlePlan）→ BattleEvent 流 → UI / 动画 / 日志")

    add_heading(doc, "2.3 依赖方向", 2)
    add_bullet(doc, "Grimhand.Battle 不依赖 UnityEngine。")
    add_bullet(doc, "Grimhand.Content 可依赖 Battle（定义数据结构）。")
    add_bullet(doc, "Grimhand.Presentation 依赖 Battle + Content。")
    add_bullet(doc, "Grimhand.Expedition / Meta 依赖 Battle，不被 Battle 反向依赖。")

    # Ch2b - confirmed rules detail
    add_heading(doc, "第二章（补充）：已确认战斗规则细则", 1)
    add_para(doc, "以下规则经项目负责人确认，优先级高于 GDD v0.4 中模糊或冲突的描述；实现与测试须严格遵循。")

    add_heading(doc, "补充 2.1 能量系统", 2)
    add_table(
        doc,
        ["规则", "说明"],
        [
            ["能量上限", "8 点（可被天赋/遗物修饰，修饰器挂在 BattleConfig）"],
            ["战斗开始", "第 1 回合开始时玩家能量 = 8（满能量开局）"],
            ["回合开始回复", "每个己方回合开始（Draw 阶段）回复 3 点能量"],
            ["溢出处理", "当前能量 + 3 后超过 8 的部分丢弃，即 min(当前+3, 8)"],
            ["跨回合保留", "上回合未消耗的能量保留至下回合，再叠加回复"],
            ["Planning 扣费", "点选加入出牌队列时立即扣除该牌费用；UI 实时显示剩余能量"],
            ["取消返还", "再次点击已选牌取消选择时，立即返还该牌费用"],
            ["不可选", "剩余能量 < 卡牌费用时，该牌不可被新选入队列"],
            ["确认出牌", "点击「出牌」仅锁定 PlayQueue 并进入 SpeedResolve，不再二次扣费"],
        ],
    )
    add_para(doc, "实现提示：Draw 阶段 EnergyRules.ApplyTurnStartRegen()；第 1 回合开局 Energy=8。Planning 使用 PlanningDraft（见补充 2.5）管理预选与能量预留。", italic=True)

    add_heading(doc, "补充 2.2 手牌、抽牌与弃牌堆（杀戮尖塔式循环）", 2)
    add_table(
        doc,
        ["规则", "说明"],
        [
            ["手牌上限", "8 张"],
            ["每回合抽牌", "Draw 阶段从抽牌堆抽 5 张"],
            ["抽牌溢出", "若抽牌后手牌 > 8，超出张数立即进入弃牌堆（不保留在手牌）"],
            ["延迟抽牌", "卡牌效果若标记「下回合抽牌」，在下一回合 Draw 阶段结算；同样受 8 张上限约束，溢出进弃牌堆"],
            ["打出", "SpeedResolve 中打出的牌进入弃牌堆"],
            ["回合结束", "EndOfTurn：手牌中所有牌进入弃牌堆"],
            ["洗回", "抽牌堆为空时仍需抽牌：将弃牌堆洗牌后作为新抽牌堆（弃牌堆清空）"],
        ],
    )
    add_para(doc, "与混合牌堆关系：仍为一个战斗级 DrawPile / DiscardPile；每张牌保留 ownerCharacterId。死亡污染规则不变。", italic=True)

    add_heading(doc, "补充 2.3 速度相同（同速）", 2)
    add_bullet(doc, "同一 SpeedResolve 轮次内，SPD 相同的角色随机排序决定先后。")
    add_bullet(doc, "随机源必须为 BattleRng(seed)，保证回放、单元测试（固定种子）可复现。")
    add_bullet(doc, "SpeedResolver 输出 ResolutionStep 时应记录本轮同速组的洗牌结果（便于 Debug）。")

    add_heading(doc, "补充 2.4 角色立绘（表现层，数据须对齐）", 2)
    add_table(
        doc,
        ["立绘状态", "触发时机"],
        [
            ["Idle", "默认；某张牌结算完毕后恢复"],
            ["Attack", "打出攻击类卡牌时，直至该牌对该角色结算完毕"],
            ["Defense", "打出防御类卡牌时，直至该牌对该角色结算完毕"],
            ["Status", "打出状态类卡牌时，直至该牌对该角色结算完毕"],
            ["Death", "角色 HP 归零后，替换 Idle"],
        ],
    )
    add_bullet(doc, "卡牌数据：CardDefinition.cardType ∈ { Attack, Defense, Status }，映射到立绘。")
    add_bullet(doc, "事件：CardPortraitPoseChanged(character, pose) 在出牌时发出；CardPortraitIdleRestored(character) 在该牌 Resolution 结束时发出。")
    add_bullet(doc, "GDD 中的施法/受伤立绘本期不做；后续可用 Status 子类型扩展。")

    add_heading(doc, "补充 2.5 出牌阶段交互（Planning UI）", 2)
    add_para(doc, "Planning 分为「预选（Draft）」与「确认提交」两阶段；能量在预选时即结算，便于玩家决策。", bold=True)

    add_heading(doc, "补充 2.5.1 PlanningDraft（预选状态）", 3)
    add_table(
        doc,
        ["操作", "行为", "能量"],
        [
            ["选牌（加入队列）", "手牌中的牌加入 PlayQueue 末尾；点选顺序=出牌顺序", "立即扣除该牌 cost"],
            ["取消选牌", "再次点击已选中的牌，从 PlayQueue 移除", "立即返还该牌 cost"],
            ["尝试选牌失败", "剩余能量 < cost，或牌不可用（死亡污染等）", "不扣费、不加入队列"],
            ["UI 显示", "始终展示「当前剩余能量」= Energy - 已选牌总费用", "实时更新"],
        ],
    )
    add_bullet(doc, "每张牌仅由其 owner 角色在 SpeedResolve 中执行；预选时 UI 须标识所属角色。")
    add_bullet(doc, "可调整：取消后重选、改变顺序（若支持拖拽排序，取消/重选逻辑同上）。")

    add_heading(doc, "补充 2.5.2 确认出牌（提交）", 3)
    add_bullet(doc, "玩家点击「出牌」→ 将 PlanningDraft 固化为 BattlePlan → 进入 SpeedResolve。")
    add_bullet(doc, "提交时仅做完整性校验（队列非空可选、牌仍合法等），不再扣除能量。")
    add_bullet(doc, "提交后直至 SpeedResolve 结束，预选队列不可再改。")

    add_heading(doc, "补充 2.5.3 架构建议", 3)
    add_bullet(doc, "BattleEngine 提供 PlanningDraft API：TrySelectCard / TryDeselectCard / GetRemainingEnergy / CommitPlan。")
    add_bullet(doc, "表现层只调用上述 API，不自行计算能量，避免 UI 与模拟状态不一致。")
    add_bullet(doc, "事件：CardSelectedForPlay, CardDeselectedFromPlay, EnergyRemainingChanged（供 UI 刷新）。")
    add_bullet(doc, "敌方 PvE：Planning 结束前展示意图；敌方计划由 IEnemyPlanner 生成，玩家不可见具体选牌。")

    # Ch3
    add_heading(doc, "第三章：战斗内核设计", 1)

    add_heading(doc, "3.1 四大核心对象", 2)
    add_table(
        doc,
        ["对象", "职责"],
        [
            ["BattleState", "可变战斗状态：Combatant、牌堆/手牌/弃牌、能量、回合数、Buff、ActionHistory"],
            ["TurnPhase FSM", "Draw → Planning → SpeedResolve → EndOfTurn 阶段机"],
            ["BattlePlan", "出牌阶段唯一输入：手牌点选顺序构成的 PlayQueue + 能量校验"],
            ["BattleEvent", "只追加事件流，供表现层/Replay/未来联网消费"],
        ],
    )

    add_heading(doc, "3.2 回合流程（对齐 GDD 2.2）", 2)
    add_table(
        doc,
        ["阶段", "行为"],
        [
            ["Draw", "回复能量（首回合开局 8，之后每回合 +3 上限 8）；抽 5 张；溢出进弃牌堆；抽牌堆空则洗回弃牌堆"],
            ["Planning", "手牌点选建 Draft（选牌即扣费、可取消返还）→ 点「出牌」提交；PvE 显示意图"],
            ["SpeedResolve", "多轮按速度交错结算（同速随机）；每轮每角色最多 1 张；打出牌入弃牌堆"],
            ["EndOfTurn", "手牌全进弃牌堆；持续效果/Buff 倒计时；判断胜负"],
        ],
    )
    add_para(doc, "PvP 与 PvE 共用阶段机；差异仅在 Planning 的输入提供者（见 3.6）。")

    add_heading(doc, "3.3 PlanningDraft 与 BattlePlan", 2)
    add_para(doc, "Planning 阶段使用 Draft；提交后生成不可变的 BattlePlan：")
    add_bullet(doc, "PlanningDraft.PlayQueue：预选的有序 List<CardInstanceId>；选牌/取消时维护。")
    add_bullet(doc, "PlanningDraft.EnergyRemaining：实时剩余能量（选牌扣费、取消返还）。")
    add_bullet(doc, "BattlePlan：由 Draft.Commit() 生成；PlayQueue 快照 + 已锁定能量消耗。")
    add_bullet(doc, "SpeedResolve 从 BattlePlan.PlayQueue 按速度轮询；同角色多张按队列中的先后。")
    add_bullet(doc, "TrySelectCard：校验 EnergyRemaining >= cost、牌在手牌、owner 存活、未污染。")
    add_bullet(doc, "CommitPlan：不再扣费；进入 SpeedResolve。")

    add_heading(doc, "3.4 SpeedResolver（速度结算模块）", 2)
    add_para(doc, "独立模块，输入双方 BattlePlan + 全体 SPD，输出 ResolutionStep 有序列表。")
    add_para(doc, "GDD 验收用例（必须写入单元测试）：", bold=True)
    add_para(
        doc,
        "玩家 A(SPD=10) 出 3 张，B(SPD=5) 出 1 张；敌人 X(SPD=9)、Y(SPD=7) 各 1 张。"
        "第一轮：A→X→Y→B；第二轮：A；第三轮：A。",
    )
    add_bullet(doc, "同速：SPD 相等时由 BattleRng 随机排序（见补充 2.3）。")
    add_bullet(doc, "扩展规则：无牌跳过、中途死亡移出队列、控制类跳过等后续增量添加。")

    add_heading(doc, "3.5 应对与反制（ActionHistory）", 2)
    add_para(doc, "维护 LastActionSnapshot：上一行动角色、行动类型（Attack/Spell/Buff/Kill 等）、目标、是否击杀。")
    add_bullet(doc, "IReactionCondition：数据驱动条件，由效果系统查询。")
    add_bullet(doc, "速度慢的角色后结算，自然更容易触发应对——无需额外硬编码平衡。")
    add_bullet(doc, "Demo 至少实现 1 种应对（建议：弹反）+ 通用条件框架。")

    add_heading(doc, "3.6 PvE / PvP 输入抽象", 2)
    add_table(
        doc,
        ["接口", "PvE", "PvP（后期）"],
        [
            ["IPlayerInputProvider", "玩家提交 BattlePlan", "双方各自提交，Planning 结束前互不可见"],
            ["IIntentProvider", "战斗前展示敌人意图", "不使用"],
            ["IEnemyPlanner", "根据意图/AI 生成敌方 BattlePlan", "由对手 BattlePlan 替代"],
        ],
    )

    add_heading(doc, "3.7 位置系统（3 槽位）", 2)
    add_table(
        doc,
        ["槽位", "效果（GDD）", "默认被攻击优先级"],
        [
            ["1 前排", "30% 减伤", "最高"],
            ["2 中排", "15% 减伤 + 15% 增伤", "中等"],
            ["3 后排", "30% 增伤", "最低"],
        ],
    )
    add_para(doc, "目标选择规则、换位效果由独立 PositionRules 模块处理。")

    add_heading(doc, "3.8 牌堆、弃牌与死亡污染", 2)
    add_bullet(doc, "30 张混合牌堆：每张牌绑定 ownerCharacterId，仅该角色可打出。")
    add_bullet(doc, "DrawPile / Hand / DiscardPile 三分；抽牌堆空时弃牌堆洗入（见第二章补充 2.2）。")
    add_bullet(doc, "手牌上限 8；每回合 Draw 抽 5；EndOfTurn 手牌全弃。")
    add_bullet(doc, "角色死亡：其牌保留在牌堆/手牌中但标记为不可使用（污染）。")
    add_bullet(doc, "Demo 可第二周实现污染，但 CardInstance 从第一天预留 IsUsable / OwnerId。")

    add_heading(doc, "3.9 胜负条件", 2)
    add_bullet(doc, "PvE 胜利：消灭所有敌人。")
    add_bullet(doc, "PvE 失败：己方全队 HP 归零（远征结束逻辑由 Expedition 层处理）。")

    # Ch4
    add_heading(doc, "第四章：卡牌与效果系统", 1)
    add_heading(doc, "4.1 设计原则", 2)
    add_bullet(doc, "禁止「每张卡一个 C# 类」；新卡 = 新 ScriptableObject + 组合 EffectSpec。")
    add_bullet(doc, "高品质卡强调机制差异，而非纯数值膨胀（对齐 GDD 5.2）。")

    add_heading(doc, "4.2 CardDefinition（ScriptableObject）", 2)
    add_bullet(doc, "ownerCharacterId, cost, cardType（Attack / Defense / Status，决定立绘）")
    add_bullet(doc, "tags（Reaction 等机制标签，与 cardType 正交）")
    add_bullet(doc, "List<EffectSpec>：可组合效果列表")

    add_heading(doc, "4.3 EffectSpec 与执行器", 2)
    add_table(
        doc,
        ["字段", "说明"],
        [
            ["trigger", "OnPlay / OnReaction / OnHit 等"],
            ["condition", "可选，如 LastAction 为 Attack 且目标为自己"],
            ["actions", "DealDamage, Heal, ApplyStatus, SwapPosition, DrawCard…"],
        ],
    )
    add_para(doc, "Demo 阶段建议实现的效果类型（约 5~8 种）：", bold=True)
    add_bullet(doc, "造成伤害、获得护甲/减伤、抽牌、反射伤害、换位、应用 Buff、回复生命。")

    # Ch5
    add_heading(doc, "第五章：事件流（BattleEvent）", 1)
    add_para(doc, "所有状态变更通过事件对外公布，表现层不得直接修改 BattleState。")
    add_heading(doc, "5.1 事件类别（示例）", 2)
    add_bullet(doc, "流程：TurnPhaseChanged, BattleStarted, BattleEnded")
    add_bullet(doc, "卡牌：CardDrawn, CardPlayed, CardResolved, CardDiscarded")
    add_bullet(doc, "战斗：DamageApplied, HealApplied, StatusApplied, ReactionTriggered")
    add_bullet(doc, "单位：CharacterDied, PositionSwapped, CardPortraitPoseChanged, CardPortraitIdleRestored")
    add_bullet(doc, "牌堆：DeckShuffled, DeckPolluted（角色死亡导致）")

    add_heading(doc, "5.2 表现层职责", 2)
    add_bullet(doc, "订阅事件队列，按序播放动画（可配置节奏）。")
    add_bullet(doc, "维护「显示用」HP/手牌，但以事件为准同步，避免双份真相。")
    add_bullet(doc, "Debug 模式可跳过动画，Instant Resolve。")

    # Ch6
    add_heading(doc, "第六章：后续系统接口（本期不实现）", 1)

    add_heading(doc, "6.1 远征（Expedition）", 2)
    add_para(doc, "ExpeditionRun 包装多场战斗，战斗引擎不感知「第几层地牢」。")
    add_bullet(doc, "输入：PartySnapshot + RunConfig（难度 Lv、RNG 种子）")
    add_bullet(doc, "输出：BattleConfig + 修饰符列表（怪物 HP/ATK 加成、临时卡、遗物）")
    add_bullet(doc, "节点三选一、商店、休息点等由 Expedition 状态机驱动。")

    add_heading(doc, "6.2 双层经验（Meta）", 2)
    add_table(
        doc,
        ["类型", "作用域", "Demo"],
        [
            ["局内经验", "当前远征", "可实现简化版"],
            ["局外经验", "永久养成", "仅定义 ISaveRepository 接口，实现可空"],
        ],
    )

    add_heading(doc, "6.3 PvP 预留", 2)
    add_bullet(doc, "BattlePlan 可序列化；全程使用种子 RNG（BattleRng(seed)）。")
    add_bullet(doc, "区分 HiddenBattleView（Planning）与 FullBattleView（结算后）。")

    add_heading(doc, "6.4 交易与拍卖行", 2)
    add_para(doc, "与战斗内核无直接耦合；属 Meta / 服务端范畴，本期不涉及。")

    # Ch7
    add_heading(doc, "第七章：PvE 最小 Demo 范围", 1)
    add_heading(doc, "7.1 必须完成（MVP）", 2)
    add_table(
        doc,
        ["项", "说明"],
        [
            ["战斗场景", "BattleSandbox 单场景 + Debug/简易 UI"],
            ["阵容", "3 玩家角色 vs 1~3 敌人"],
            ["卡组", "每人 10 张混合牌堆；手牌上限 8；每回合抽 5；能量上限 8、首回合满能量、每回合 +3"],
            ["出牌 UI", "选牌即扣费/可取消/能量不足不可选 + 出牌按钮仅提交"],
            ["立绘", "Attack/Defense/Status + Idle/Death 占位切换"],
            ["核心流程", "Draw → Planning → SpeedResolve → EndOfTurn"],
            ["速度结算", "通过 GDD 示例单元测试"],
            ["意图", "至少 1 种敌人意图（攻击/防御）"],
            ["应对", "至少 1 种（弹反）"],
            ["位置", "3 槽位减伤/增伤 + 前排优先承伤"],
            ["效果", "5~8 种 EffectSpec"],
        ],
    )

    add_heading(doc, "7.2 明确不做（留接口）", 2)
    add_bullet(doc, "远征三选一地图、遗物、消耗品、商店")
    add_bullet(doc, "局外天赋、装备系统（可用写死属性代替）")
    add_bullet(doc, "PvP 盲选、拍卖行")
    add_bullet(doc, "正式美术立绘资源（Demo 用占位图 + CardPortrait 事件即可）")
    add_bullet(doc, "死亡污染可列为 Demo 第二期，但数据结构须预留")

    add_heading(doc, "7.3 Demo 成功标准", 2)
    add_bullet(doc, "单元测试：GDD 速度结算示例 100% 通过；同速随机用固定种子可复现。")
    add_bullet(doc, "单元测试：能量保留+回复、手牌上限溢出、弃牌堆洗回。")
    add_bullet(doc, "可玩：一场 3v1 或 3v3 战斗从开始到胜负可重复完成。")
    add_bullet(doc, "架构：无 MonoBehaviour 内写伤害公式；无 UI 直接改 BattleState。")
    add_bullet(doc, "扩展：新增 1 张卡仅需新增 ScriptableObject，不改引擎代码。")

    # Ch8
    add_heading(doc, "第八章：目录与程序集结构", 1)
    add_para(doc, "建议 Unity 目录：")
    add_para(
        doc,
        "Assets/_Project/Scripts/Core/\n"
        "Assets/_Project/Scripts/Battle/Engine/\n"
        "Assets/_Project/Scripts/Battle/Rules/\n"
        "Assets/_Project/Scripts/Battle/Reactions/\n"
        "Assets/_Project/Scripts/Battle/Effects/\n"
        "Assets/_Project/Scripts/Battle/Events/\n"
        "Assets/_Project/Scripts/Content/\n"
        "Assets/_Project/Scripts/Presentation/Battle/\n"
        "Assets/_Project/Scripts/Presentation/UI/\n"
        "Assets/_Project/Scripts/Expedition/（空或接口）\n"
        "Assets/_Project/Scripts/Meta/\n"
        "Assets/_Project/Data/Cards|Characters|Enemies/\n"
        "Assets/_Project/Scenes/BattleSandbox.unity\n"
        "Assets/_Project/Tests/Battle/",
    )

    # Ch9
    add_heading(doc, "第九章：开发顺序与里程碑", 1)
    add_table(
        doc,
        ["阶段", "目标", "预估"],
        [
            ["0", "纸面规则与 GDD 对齐，确认枚举与术语", "0.5 天"],
            ["1", "Battle 纯逻辑 + SpeedResolver 单元测试", "2~3 天"],
            ["2", "Effect 执行器 + 5~8 种效果", "2~3 天"],
            ["3", "Planning + SpeedResolve 端到端", "2 天"],
            ["4", "BattleSandbox + Debug UI", "2~3 天"],
            ["5", "敌人意图 + 1 种应对 + 位置规则", "3~4 天"],
            ["6", "死亡污染 / 换位（可选第二期）", "2 天"],
            ["7", "Expedition 壳 + 线性 3 场战斗", "后续迭代"],
        ],
    )
    add_para(doc, "第 1 周里程碑：测试跑通速度示例 + 手动完成一场 3v1。", bold=True)

    # Ch10
    add_heading(doc, "第十章：测试策略", 1)
    add_bullet(doc, "SpeedResolver、EnergyRules、ReactionConditions 必须有纯 C# 单元测试。")
    add_bullet(doc, "效果组合测试：用最小 CardDefinition 夹具验证伤害/护甲/应对链。")
    add_bullet(doc, "可选：Golden Test——固定种子完整战斗日志比对。")
    add_bullet(doc, "Play Mode 仅做 UI 冒烟，核心规则不依赖 Enter Play Mode。")

    # Ch11
    add_heading(doc, "第十一章：反模式（禁止）", 1)
    add_table(
        doc,
        ["反模式", "后果", "替代方案"],
        [
            ["MonoBehaviour 内写战斗公式", "不可测、难做 PvP", "BattleEngine"],
            ["每张卡一个 XXXCard.cs", "300 张卡维护崩溃", "EffectSpec + SO"],
            ["UI 直接改 HP/手牌", "状态不一致", "消费 BattleEvent"],
            ["先做地图再做战斗", "核心未验证", "按 GDD 第十章顺序"],
            ["全局 Singleton 管一切", "耦合、难测", "BattleSession 实例"],
        ],
    )

    # Ch12
    add_heading(doc, "第十二章：技术栈与现有工程", 1)
    add_bullet(doc, "引擎：Unity（当前为 URP 2D 模板）")
    add_bullet(doc, "已具备：Input System、2D Animation、Physics2D")
    add_bullet(doc, "Demo 不强制引入：Addressables、Netcode（内容量/PvP 阶段再评估）")
    add_bullet(doc, "立绘：cardType → Attack/Defense/Status 立绘；CardPortraitPoseChanged / IdleRestored 事件驱动")

    # Ch13
    add_heading(doc, "第十三章：审核检查清单", 1)
    add_para(doc, "请审核人逐项确认：")
    add_bullet(doc, "□ 分层与程序集划分是否可执行？")
    add_bullet(doc, "□ 战斗内核四对象是否覆盖 GDD 核心机制？")
    add_bullet(doc, "□ Demo 范围是否足够小且可验证？")
    add_bullet(doc, "□ 事件流方案是否满足未来 PvP/Replay？")
    add_bullet(doc, "□ EffectSpec 数据驱动是否满足卡牌规模？")
    add_bullet(doc, "□ 开发顺序与里程碑是否合理？")
    add_bullet(doc, "□ v1.1/v1.2 已确认规则是否已全部纳入第二章补充？")
    add_bullet(doc, "□ Planning 选牌即扣费、取消返还是否符合预期？")
    add_bullet(doc, "□ 是否有遗漏的 GDD 机制需在 Demo 前纳入内核？")
    add_para(doc, "审核结论：□ 通过，可启动 Demo  □ 需修订（意见：________________________）")

    add_heading(doc, "附录 A：与 GDD v0.4 章节对照", 1)
    add_table(
        doc,
        ["GDD 章节", "本框架对应"],
        [
            ["第二章 核心战斗", "第三、四、五章"],
            ["第三章 Roguelike 远征", "第六章 6.1（接口）"],
            ["第四章 角色养成", "第六章 6.2（接口）"],
            ["第五章 卡牌系统", "第四章"],
            ["第八章 PvP", "第三章 3.6、第六章 6.3"],
            ["第九章 美术立绘", "第二章补充 2.4 + CardPortrait 事件（简化为三类+Idle/Death）"],
            ["v1.1/v1.2 审核修订", "第一章 1.4 + 第二章补充（含 PlanningDraft 扣费）"],
            ["第十章 下一步行动", "第七、九章"],
        ],
    )

    doc.add_paragraph()
    add_para(doc, "—— 文档结束 ——", bold=True)
    add_para(doc, "生成工具：Grimhand/Tools/generate_framework_doc.py（可重复生成修订版）", italic=True)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
